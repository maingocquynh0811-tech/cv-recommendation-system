"""
Module xử lý input từ người dùng (text query + file PDF/Word)
"""

import re
import fitz  # PyMuPDF
from docx import Document
from typing import Dict, Optional, Tuple
from utils.logger import get_logger

logger = get_logger(__name__)

class InputProcessor:
    def __init__(self, max_file_size_mb=50, max_pages=100):
        """
        Khởi tạo input processor
        
        Args:
            max_file_size_mb: Kích thước file tối đa (MB)
            max_pages: Số trang tối đa cho PDF
        """
        self.max_file_size = max_file_size_mb * 1024 * 1024  # Convert to bytes
        self.max_pages = max_pages
    
    def clean_text(self, text: str) -> str:
        """
        Làm sạch text
        
        Args:
            text: Raw text
            
        Returns:
            Cleaned text
        """
        if not text:
            return ""
        
        # Xóa ký tự đặc biệt và emoji
        text = re.sub(r'[^\w\s\-.,!?():;]', ' ', text)
        
        # Chuẩn hóa khoảng trắng
        text = ' '.join(text.split())
        
        return text
    
    def extract_pdf_text(self, pdf_path: str) -> Tuple[Optional[str], Optional[Dict]]:
    
        doc = None
        try:
            import os
            
            # Kiểm tra file có tồn tại không
            if not os.path.exists(pdf_path):
                logger.error(f"PDF file not found: {pdf_path}")
                return None, None
            
            # Kiểm tra file size
            file_size = os.path.getsize(pdf_path)
            if file_size > self.max_file_size:
                logger.error(f"PDF file too large: {file_size / 1024 / 1024:.2f} MB")
                return None, None
            
            # Mở PDF
            doc = fitz.open(pdf_path)
            
            # Kiểm tra số trang
            page_count = len(doc)
            if page_count > self.max_pages:
                logger.warning(f"PDF has {page_count} pages, limiting to {self.max_pages}")
            
            # Trích xuất metadata
            metadata = {
                'page_count': page_count,
                'title': doc.metadata.get('title', '') if doc.metadata else '',
                'author': doc.metadata.get('author', '') if doc.metadata else '',
                'keywords': doc.metadata.get('keywords', '') if doc.metadata else ''
            }
            
            # Trích xuất text từ mỗi trang
            all_text = []
            pages_to_process = min(page_count, self.max_pages)
            
            for page_num in range(pages_to_process):
                try:
                    page = doc[page_num]
                    text = page.get_text()
                    if text and text.strip():
                        all_text.append(text)
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num}: {e}")
                    continue
            
            # Ghép text
            full_text = '\n'.join(all_text)
            
            if not full_text or not full_text.strip():
                logger.error("No text extracted from PDF - might be a scanned image")
                return None, None
            
            logger.info(f"Extracted {len(full_text)} characters from PDF ({page_count} pages)")
            
            return full_text, metadata
            
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            return None, None
        
        finally:
            # Đảm bảo đóng document
            if doc is not None:
                try:
                    doc.close()
                except:
                    pass
    
    def extract_word_text(self, docx_path: str) -> Tuple[Optional[str], Optional[Dict]]:
        """
        Trích xuất text từ Word document
        
        Args:
            docx_path: Đường dẫn đến file Word
            
        Returns:
            Tuple (extracted_text, metadata)
        """
        try:
            import os
            
            # Kiểm tra file size
            file_size = os.path.getsize(docx_path)
            if file_size > self.max_file_size:
                logger.error(f"Word file too large: {file_size / 1024 / 1024:.2f} MB")
                return None, None
            
            # Mở document
            doc = Document(docx_path)
            
            # Trích xuất metadata
            core_props = doc.core_properties
            metadata = {
                'paragraph_count': len(doc.paragraphs),
                'title': core_props.title or '',
                'author': core_props.author or '',
                'keywords': core_props.keywords or ''
            }
            
            # Trích xuất text từ paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # Trích xuất text từ tables (optional)
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' '.join(cell.text for cell in row.cells)
                    if row_text.strip():
                        table_texts.append(row_text)
            
            # Ghép text
            full_text = '\n'.join(paragraphs)
            if table_texts:
                full_text += '\n' + '\n'.join(table_texts)
            
            logger.info(f"Extracted {len(full_text)} characters from Word ({len(paragraphs)} paragraphs)")
            return full_text, metadata
            
        except Exception as e:
            logger.error(f"Error extracting Word text: {e}")
            return None, None
    
    def extract_structured_info(self, text: str) -> Dict:
        """
        Trích xuất thông tin có cấu trúc từ text (Abstract, Introduction, etc.)
        
        Args:
            text: Full text từ document
            
        Returns:
            Dictionary chứa các sections
        """
        sections = {
            'abstract': '',
            'introduction': '',
            'conclusion': '',
            'keywords': [],
            'section_headings': []
        }
        
        # Tách thành dòng
        lines = text.split('\n')
        
        # Tìm Abstract
        abstract_pattern = r'(?i)^(abstract|summary)[:\s]*$'
        for i, line in enumerate(lines):
            if re.match(abstract_pattern, line.strip()):
                # Lấy 10 dòng tiếp theo
                abstract_lines = []
                for j in range(i+1, min(i+15, len(lines))):
                    if lines[j].strip() and not re.match(r'(?i)^(introduction|1\.|keywords)', lines[j].strip()):
                        abstract_lines.append(lines[j].strip())
                    else:
                        break
                sections['abstract'] = ' '.join(abstract_lines)
                break
        
        # Tìm Introduction
        intro_pattern = r'(?i)^(introduction|1\.|1\s+introduction)[:\s]*$'
        for i, line in enumerate(lines):
            if re.match(intro_pattern, line.strip()):
                intro_lines = []
                for j in range(i+1, min(i+20, len(lines))):
                    if lines[j].strip() and not re.match(r'(?i)^(2\.|methodology|method)', lines[j].strip()):
                        intro_lines.append(lines[j].strip())
                    else:
                        break
                sections['introduction'] = ' '.join(intro_lines)
                break
        
        # Tìm Keywords
        keyword_pattern = r'(?i)^(keywords|key words)[:\s]*(.+)$'
        for line in lines:
            match = re.match(keyword_pattern, line.strip())
            if match:
                keywords_text = match.group(2)
                sections['keywords'] = [k.strip() for k in re.split(r'[,;]', keywords_text) if k.strip()]
                break
        
        # Tìm Section headings (các heading có số hoặc chữ in hoa)
        heading_pattern = r'^(\d+\.|\d+\s+|[A-Z][A-Z\s]+)(.{3,50})$'
        for line in lines:
            line = line.strip()
            if re.match(heading_pattern, line) and len(line) < 100:
                sections['section_headings'].append(line)
        
        return sections
    
    def preprocess_file_text(self, raw_text: str, metadata: Dict) -> str:
        """
        Xử lý và làm sạch text từ file
        
        Args:
            raw_text: Raw text từ file
            metadata: Metadata từ file
            
        Returns:
            Processed text
        """
        # Làm sạch text
        text = self.clean_text(raw_text)
        
        # Trích xuất thông tin cấu trúc
        structured = self.extract_structured_info(text)
        
        # Xây dựng văn bản đại diện theo chiến lược weighted combination
        # Title (×3) + Keywords (×2) + Abstract (×2) + Introduction + Section Headings (×2)
        
        processed_parts = []
        
        # Title
        title = metadata.get('title', '').strip()
        if title:
            processed_parts.extend([title] * 3)
        
        # Keywords
        keywords = structured.get('keywords', [])
        if keywords:
            keywords_text = ' '.join(keywords)
            processed_parts.extend([keywords_text] * 2)
        
        # Abstract
        abstract = structured.get('abstract', '').strip()
        if abstract:
            processed_parts.extend([abstract] * 2)
        
        # Introduction
        introduction = structured.get('introduction', '').strip()
        if introduction:
            processed_parts.append(introduction)
        
        # Section Headings
        headings = structured.get('section_headings', [])
        if headings:
            headings_text = ' '.join(headings[:10])  # Top 10 headings
            processed_parts.extend([headings_text] * 2)
        
        # Nếu không có thông tin cấu trúc, dùng 2000 ký tự đầu
        if not processed_parts:
            processed_parts = [text[:2000]]
        
        # Ghép lại
        final_text = ' '.join(processed_parts)
        
        # Giới hạn độ dài (max ~4000 tokens = 16000 chars)
        if len(final_text) > 16000:
            final_text = final_text[:16000]
        
        logger.info(f"Processed file text: {len(final_text)} characters")
        return final_text
    
    def process_query_text(self, query: str) -> str:
        """
        Xử lý text query từ user
        
        Args:
            query: Raw query text
            
        Returns:
            Cleaned query
        """
        # Làm sạch
        query = self.clean_text(query)
        
        # Trim
        query = query.strip()
        
        logger.info(f"Processed query: '{query}' ({len(query)} characters)")
        return query
    
    def combine_query_and_file(self, query_text: str, file_text: str, strategy='weighted') -> str:
        """
        Kết hợp text query và file content
        
        Args:
            query_text: Text từ user query
            file_text: Text từ file
            strategy: Chiến lược kết hợp ('simple', 'weighted')
            
        Returns:
            Combined text
        """
        if strategy == 'simple':
            # Simple concatenation
            combined = f"{query_text}\n\n{file_text}"
        
        elif strategy == 'weighted':
            # Weighted: query có trọng số cao hơn (×3)
            combined = f"{query_text} {query_text} {query_text}\n\n{file_text}"
        
        else:
            combined = f"{query_text}\n\n{file_text}"
        
        # Giới hạn độ dài
        if len(combined) > 20000:
            combined = combined[:20000]
        
        logger.info(f"Combined text: {len(combined)} characters")
        return combined
    
    def process_input(self, query_text: Optional[str] = None, 
                      file_path: Optional[str] = None,
                      file_type: Optional[str] = None) -> Optional[str]:
        """
        Xử lý toàn bộ input từ user
        
        Args:
            query_text: Text query (optional)
            file_path: Đường dẫn file (optional)
            file_type: Loại file ('pdf' hoặc 'word')
            
        Returns:
            Processed text ready for embedding
        """
        final_text = ""
        
        # Case 1: Chỉ có query text
        if query_text and not file_path:
            logger.info("Processing query text only")
            final_text = self.process_query_text(query_text)
        
        # Case 2: Chỉ có file
        elif file_path and not query_text:
            logger.info(f"Processing file only: {file_path}")
            
            if file_type == 'pdf':
                raw_text, metadata = self.extract_pdf_text(file_path)
            elif file_type == 'word':
                raw_text, metadata = self.extract_word_text(file_path)
            else:
                logger.error(f"Unsupported file type: {file_type}")
                return None
            
            if raw_text:
                final_text = self.preprocess_file_text(raw_text, metadata)
            else:
                return None
        
        # Case 3: Có cả query và file
        elif query_text and file_path:
            logger.info(f"Processing query + file: {file_path}")
            
            # Xử lý query
            processed_query = self.process_query_text(query_text)
            
            # Xử lý file
            if file_type == 'pdf':
                raw_text, metadata = self.extract_pdf_text(file_path)
            elif file_type == 'word':
                raw_text, metadata = self.extract_word_text(file_path)
            else:
                logger.error(f"Unsupported file type: {file_type}")
                return None
            
            if raw_text:
                processed_file = self.preprocess_file_text(raw_text, metadata)
                final_text = self.combine_query_and_file(processed_query, processed_file, strategy='weighted')
            else:
                # Nếu file lỗi, dùng query
                final_text = processed_query
        
        else:
            logger.error("No input provided (no query and no file)")
            return None
        
        if not final_text or not final_text.strip():
            logger.error("Failed to process input: empty result")
            return None
        
        return final_text